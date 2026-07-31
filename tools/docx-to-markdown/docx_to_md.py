"""
Convert Word (.docx) files to plain Markdown text.

Usage:
    python docx_to_md.py path/to/file.docx
    python docx_to_md.py path/to/folder

If a folder is given, every .docx inside it is converted.
Each .docx gets a matching .md file written next to it.

Note: only modern .docx is supported, not the old binary .doc format.
"""

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

HEADING_STYLES = {
    "Title": "# ",
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
}


def paragraph_to_md(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    style = paragraph.style.name if paragraph.style else ""
    if style in HEADING_STYLES:
        return HEADING_STYLES[style] + text
    if style.startswith("List Bullet"):
        return "- " + text
    if style.startswith("List Number"):
        return "1. " + text
    return text


def table_to_md(table) -> str:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def convert_docx(docx_path: Path) -> Path:
    md_path = docx_path.with_suffix(".md")
    document = Document(docx_path)

    paragraphs_by_elem = {p._p: p for p in document.paragraphs}
    tables_by_elem = {t._tbl: t for t in document.tables}

    blocks = []
    for kind, elem in iter_block_items(document):
        if kind == "paragraph":
            text = paragraph_to_md(paragraphs_by_elem[elem])
            if text:
                blocks.append(text)
        else:
            md_table = table_to_md(tables_by_elem[elem])
            if md_table:
                blocks.append(md_table)

    md_path.write_text("\n\n".join(blocks), encoding="utf-8")
    return md_path


def main() -> None:
    if len(sys.argv) != 2:
        print("Usage: python docx_to_md.py <file.docx | folder>")
        sys.exit(1)

    target = Path(sys.argv[1])

    if target.is_dir():
        docs = sorted(target.glob("*.docx"))
        if not docs:
            print(f"No .docx files found in {target}")
            sys.exit(1)
    elif target.is_file() and target.suffix.lower() == ".docx":
        docs = [target]
    else:
        print(f"Not a .docx file or folder: {target}")
        sys.exit(1)

    for docx_path in docs:
        md_path = convert_docx(docx_path)
        print(f"Converted: {docx_path.name} -> {md_path.name}")


if __name__ == "__main__":
    main()
